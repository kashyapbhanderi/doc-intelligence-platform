"""
scripts/create_test_docs.py
Creates sample test files for editor tool testing.
Run once to generate: data/test_docs/
"""
import os
from pathlib import Path

OUTPUT_DIR = Path("data/test_docs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def create_test_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(
        "AI Research Summary Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro paragraph
    doc.add_paragraph(
        "This document summarises key findings from "
        "recent AI/ML research papers. "
        "The retrieval augmented generation (RAG) "
        "approach has shown significant improvements "
        "over standard fine-tuning methods."
    )

    # Section heading
    doc.add_heading("Key Findings", level=2)

    # Bullet points
    for item in [
        "RAG reduces hallucination by 40% vs baseline",
        "LoRA fine-tuning requires 90% less GPU memory",
        "Hybrid search outperforms BM25 by 15% NDCG",
        "Multi-agent systems improve answer faithfulness",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Table
    doc.add_heading("Performance Metrics", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Method"
    header[1].text = "NDCG@10"
    header[2].text = "Latency (s)"

    rows = [
        ("Baseline BM25",   "0.54", "1.2"),
        ("Vector Search",   "0.61", "1.4"),
        ("Hybrid RRF",      "0.73", "1.5"),
        ("Fine-tuned RAG",  "0.81", "2.1"),
    ]
    for method, ndcg, latency in rows:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = ndcg
        row[2].text = latency

    doc.add_paragraph()
    doc.add_paragraph(
        "Contact: kashyap@example.com | "
        "Version: DRAFT | Date: 2025-01-01"
    )

    path = OUTPUT_DIR / "test_report.docx"
    doc.save(path)
    print(f"Created: {path}")
    return str(path)


def create_test_pdf():
    """Create a simple test PDF using reportlab if available,
    or copy an existing PDF from data/raw."""
    import shutil
    raw_dir = Path("data/raw")
    pdfs = list(raw_dir.glob("*.pdf"))
    if pdfs:
        dest = OUTPUT_DIR / "test_document.pdf"
        shutil.copy(pdfs[0], dest)
        print(f"Copied: {dest} (from {pdfs[0].name})")
        return str(dest)
    else:
        print("No PDFs in data/raw — skipping PDF test file")
        return None


def create_test_image():
    from PIL import Image, ImageDraw, ImageFont
    img  = Image.new("RGB", (800, 400),
                     color=(240, 248, 255))
    draw = ImageDraw.Draw(img)

    # Background grid
    for x in range(0, 800, 50):
        draw.line([(x, 0), (x, 400)],
                  fill=(200, 220, 240), width=1)
    for y in range(0, 400, 50):
        draw.line([(0, y), (800, y)],
                  fill=(200, 220, 240), width=1)

    # Main text
    draw.rectangle([(150, 120), (650, 280)],
                   fill=(255, 255, 255),
                   outline=(100, 149, 237), width=3)
    draw.text((400, 180),
              "Doc Intelligence Platform",
              fill=(30, 30, 100),
              anchor="mm")
    draw.text((400, 230),
              "Sample Test Image — Day 19",
              fill=(80, 80, 80),
              anchor="mm")

    path = OUTPUT_DIR / "test_image.png"
    img.save(path)
    print(f"Created: {path}")
    return str(path)


if __name__ == "__main__":
    print("Creating test documents...")
    print("=" * 40)
    create_test_docx()
    create_test_pdf()
    create_test_image()
    print("\nAll test files ready in data/test_docs/")