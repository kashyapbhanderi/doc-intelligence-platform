"""
agents/editor_tools/docx_tools.py
Word document editing tools.

All functions follow the same pattern:
  - Take a file path + parameters
  - Edit the document
  - Save it (in-place or to new path)
  - Return a human-readable result string

This return string is what the Editor Agent
reports back to the user.
"""
import os
import sys
sys.path.insert(0, os.path.abspath('.'))

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ── Core editing functions ────────────────────────────────

def edit_docx_text(
    path: str,
    old_text: str,
    new_text: str
) -> str:
    """
    Find and replace text in a Word document.

    Searches all paragraphs and table cells.
    Preserves original formatting of the run.

    Industry use: updating contract terms,
    changing names/dates across a document,
    template population.

    Args:
        path:     Path to .docx file
        old_text: Text to find
        new_text: Replacement text

    Returns:
        Result message with replacement count
    """
    if not os.path.exists(path):
        return f"Error: file not found: {path}"

    doc     = Document(path)
    count   = 0

    # Search paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(
                    old_text, new_text)
                count += 1

    # Search table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(
                                old_text, new_text)
                            count += 1

    doc.save(path)
    return (f"Replaced '{old_text}' → '{new_text}' "
            f"({count} occurrence(s)) in {Path(path).name}")


def insert_table_docx(
    path: str,
    headers: list,
    rows: list,
    heading: str = ""
) -> str:
    """
    Insert a formatted table into a Word document.

    Appends table at the end of the document.
    Uses 'Table Grid' style for clean borders.

    Args:
        path:    Path to .docx file
        headers: List of column header strings
        rows:    List of lists (each inner = one row)
        heading: Optional heading to add above table

    Returns:
        Result message
    """
    if not os.path.exists(path):
        return f"Error: file not found: {path}"

    doc = Document(path)

    if heading:
        doc.add_heading(heading, level=2)

    # Create table
    table = doc.add_table(
        rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row — bold
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = header_cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < len(row):
                row[i].text = str(val)

    doc.save(path)
    return (f"Inserted table ({len(headers)} cols × "
            f"{len(rows)} rows) into "
            f"{Path(path).name}")


def add_heading_docx(
    path: str,
    heading_text: str,
    level: int = 1
) -> str:
    """
    Add a heading at the end of a Word document.

    Args:
        path:         Path to .docx file
        heading_text: The heading text
        level:        1 = H1, 2 = H2, 3 = H3

    Returns:
        Result message
    """
    if not os.path.exists(path):
        return f"Error: file not found: {path}"

    doc = Document(path)
    doc.add_heading(heading_text,
                    level=max(1, min(level, 9)))
    doc.save(path)
    return (f"Added H{level} heading: "
            f"'{heading_text}' to "
            f"{Path(path).name}")


def read_docx_text(path: str) -> str:
    """
    Extract all text from a Word document.
    Used by the Editor Agent to read document
    content before making edits.

    Returns:
        Full document text as a string
    """
    if not os.path.exists(path):
        return f"Error: file not found: {path}"

    doc   = Document(path)
    lines = []

    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)

    return "\n".join(lines)


def get_docx_info(path: str) -> dict:
    """
    Get metadata about a Word document.

    Returns:
        Dict with paragraph count, table count,
        word count, and first 200 chars preview.
    """
    if not os.path.exists(path):
        return {"error": f"File not found: {path}"}

    doc        = Document(path)
    full_text  = read_docx_text(path)
    word_count = len(full_text.split())

    return {
        "filename":        Path(path).name,
        "paragraphs":      len(doc.paragraphs),
        "tables":          len(doc.tables),
        "word_count":      word_count,
        "preview":         full_text[:200],
    }


if __name__ == "__main__":
    TEST_PATH = "data/test_docs/test_report.docx"

    if not os.path.exists(TEST_PATH):
        print("Test file missing. Run:")
        print("  python scripts/create_test_docs.py")
    else:
        print("Testing DOCX Tools")
        print("=" * 50)

        # Test 1: read info
        info = get_docx_info(TEST_PATH)
        print(f"File info:")
        for k, v in info.items():
            if k != "preview":
                print(f"  {k}: {v}")

        # Test 2: find and replace
        result = edit_docx_text(
            TEST_PATH, "DRAFT", "FINAL")
        print(f"\nEdit: {result}")

        # Test 3: insert table
        result = insert_table_docx(
            TEST_PATH,
            headers=["Agent", "Role", "Status"],
            rows=[
                ["Planner",  "Query decomposition", "✅"],
                ["Executor", "Retrieval + answer",  "✅"],
                ["Critic",   "Faithfulness check",  "✅"],
            ],
            heading="Agent Pipeline Status"
        )
        print(f"Table: {result}")

        # Test 4: add heading
        result = add_heading_docx(
            TEST_PATH,
            "Appendix: Technical Details",
            level=2
        )
        print(f"Heading: {result}")

        print("\n✅ All DOCX tools working!")
        print(f"Check the output: {TEST_PATH}")